from docx import Document
from difflib import HtmlDiff

doc1 = Document("old.docx")
doc2 = Document("new.docx")

text1 = "\n".join(p.text for p in doc1.paragraphs)
text2 = "\n".join(p.text for p in doc2.paragraphs)

diff = HtmlDiff().make_file(text1.splitlines(), text2.splitlines())
with open("diff.html", "w", encoding="utf-8") as f:
    f.write(diff)